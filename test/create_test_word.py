#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
创建测试Word文件
"""

import os
from docx import Document

def create_test_docx():
    """创建测试.docx文件"""
    # 确保目录存在
    os.makedirs("uploads/word", exist_ok=True)
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档', 0)
    
    # 添加段落
    doc.add_paragraph('这是一个测试文档，用于验证Word文件处理器是否正常工作。')
    
    doc.add_paragraph('文档包含以下内容：')
    
    # 添加列表
    doc.add_paragraph('• 标题和段落', style='List Bullet')
    doc.add_paragraph('• 中文文本', style='List Bullet')
    doc.add_paragraph('• 特殊字符：@#$%^&*()', style='List Bullet')
    doc.add_paragraph('• 数字：123456789', style='List Bullet')
    
    # 添加更多段落
    doc.add_paragraph('这是另一个段落，包含更多的文本内容。这个段落应该足够长，以便测试文本分块功能。')
    
    doc.add_paragraph('最后一段，用于测试文档的完整性。')
    
    # 保存文件
    file_path = "uploads/word/test.docx"
    doc.save(file_path)
    print(f"测试.docx文件已创建: {file_path}")
    print(f"文件大小: {os.path.getsize(file_path)} 字节")
    
    return file_path

if __name__ == "__main__":
    create_test_docx() 