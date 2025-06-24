from .base_processor import BaseFileProcessor
import logging
from typing import List, Dict, Any
import docx
import os

logger = logging.getLogger(__name__)

class WordProcessor(BaseFileProcessor):
    """Word文件处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        
    def can_process(self, file_path: str) -> bool:
        """检查是否可以处理该文件"""
        return any(file_path.lower().endswith(ext) for ext in self.supported_extensions)
        
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """处理Word文件，正文和表格分别分块"""
        try:
            # 获取文件元数据
            metadata = self.get_file_metadata(file_path)
            all_chunks = []
            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext == '.docx':
                doc = docx.Document(file_path)
                # 1. 处理正文段落
                para_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
                if para_text.strip():
                    para_chunks = self.split_text(para_text)
                    for chunk in para_chunks:
                        all_chunks.append(self.create_chunk(chunk, metadata))
                # 2. 每个表格单独为一个文本块
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_rows.append(row_text)
                    table_text = "\n".join(table_rows)
                    if table_text.strip():
                        table_metadata = metadata.copy()
                        table_metadata["type"] = "table"
                        table_metadata["table_index"] = i
                        all_chunks.append(self.create_chunk(table_text, table_metadata))
            elif file_ext == '.doc':
                content = self._read_doc(file_path)
                if content.strip():
                    chunks = self.split_text(content)
                    for chunk in chunks:
                        all_chunks.append(self.create_chunk(chunk, metadata))
            else:
                raise ValueError(f"Unsupported file extension: {file_ext}")
            return all_chunks
        except Exception as e:
            logger.error(f"处理Word文件失败: {str(e)}")
            return []
    
    def _read_docx(self, file_path: str) -> str:
        """读取.docx文件，包括段落和表格"""
        try:
            doc = docx.Document(file_path)
            content = ""
            # 先提取段落
            for para in doc.paragraphs:
                content += para.text + "\n"
            # 再提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    content += row_text + "\n"
            return content
        except Exception as e:
            logger.error(f"读取.docx文件失败: {str(e)}")
            raise
    
    def _convert_doc_to_docx_with_wps(self, doc_path: str) -> str:
        """
        使用WPS自动将.doc转换为.docx，返回新文件路径，失败返回None
        """
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            wps = win32com.client.Dispatch("Kwps.Application")
            wps.Visible = False
            docx_path = doc_path + "x"  # 直接加x变成.docx
            doc = wps.Documents.Open(os.path.abspath(doc_path))
            doc.SaveAs(os.path.abspath(docx_path), FileFormat=12)  # 12 = docx
            doc.Close()
            wps.Quit()
            if os.path.exists(docx_path):
                return docx_path
            return None
        except Exception as e:
            logger.warning(f"WPS自动转换.doc为.docx失败: {str(e)}")
            return None

    def _read_doc(self, file_path: str) -> str:
        """读取.doc文件"""
        try:
            # 优先尝试用WPS自动转换为docx
            docx_path = self._convert_doc_to_docx_with_wps(file_path)
            if docx_path and os.path.exists(docx_path):
                logger.info(f"WPS自动转换.doc为.docx成功: {docx_path}")
                return self._read_docx(docx_path)
            # 首先尝试使用python-docx2txt库（主要用于.docx，但有时也能处理.doc）
            try:
                import docx2txt
                content = docx2txt.process(file_path)
                if content and content.strip():
                    logger.info("使用docx2txt成功读取.doc文件")
                    return content
            except Exception as e:
                logger.warning(f"docx2txt读取.doc文件失败: {str(e)}")
            # 尝试使用olefile读取.doc文件
            try:
                import olefile
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFileIO(file_path)
                    # 尝试读取WordDocument流
                    if ole.exists('WordDocument'):
                        word_doc = ole.openstream('WordDocument').read()
                        # 尝试读取内容
                        content = self._extract_text_from_ole(ole)
                        if content and content.strip():
                            logger.info("使用olefile成功读取.doc文件")
                            return content
                    ole.close()
            except ImportError:
                logger.warning("olefile库未安装")
            except Exception as e:
                logger.warning(f"olefile读取.doc文件失败: {str(e)}")
            # 尝试使用mammoth库
            try:
                import mammoth
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    if result.value and result.value.strip():
                        logger.info("使用mammoth成功读取.doc文件")
                        return result.value
            except ImportError:
                logger.warning("mammoth库未安装")
            except Exception as e:
                logger.warning(f"mammoth读取.doc文件失败: {str(e)}")
            # 尝试使用antiword（Linux/Unix工具）
            try:
                import subprocess
                result = subprocess.run(['antiword', file_path], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    logger.info("使用antiword成功读取.doc文件")
                    return result.stdout
            except (ImportError, FileNotFoundError, subprocess.TimeoutExpired) as e:
                logger.warning(f"antiword未安装或执行失败: {str(e)}")
            # 尝试使用catdoc（Linux/Unix工具）
            try:
                import subprocess
                result = subprocess.run(['catdoc', file_path], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    logger.info("使用catdoc成功读取.doc文件")
                    return result.stdout
            except (ImportError, FileNotFoundError, subprocess.TimeoutExpired) as e:
                logger.warning(f"catdoc未安装或执行失败: {str(e)}")
            # 如果所有方法都失败，抛出详细的异常信息
            error_msg = (
                "无法读取.doc文件。请尝试以下解决方案：\n"
                "1. 将.doc文件转换为.docx格式（推荐）：\n"
                "   - 用Microsoft Word或WPS打开文件，另存为.docx格式\n"
                "   - 或使用在线转换工具：https://convertio.co/doc-docx/\n"
                "2. 在Windows系统上安装LibreOffice，然后使用命令行转换：\n"
                "   soffice --headless --convert-to docx:MS Word 2007 XML 文件路径.doc\n"
                "3. 使用WPS Office或其他支持.doc的办公软件打开并另存为.docx\n"
                "4. 如果文件很重要，建议使用Microsoft Word进行转换"
            )
            raise Exception(error_msg)
        except Exception as e:
            logger.error(f"读取.doc文件失败: {str(e)}")
            raise
    
    def _extract_text_from_ole(self, ole):
        """从OLE文件中提取文本"""
        try:
            content = ""
            
            # 尝试读取不同的流
            streams_to_try = ['WordDocument', 'Contents', '1Table', '0Table']
            
            for stream_name in streams_to_try:
                if ole.exists(stream_name):
                    try:
                        stream = ole.openstream(stream_name)
                        data = stream.read()
                        # 尝试解码为文本
                        try:
                            text = data.decode('utf-8', errors='ignore')
                            # 清理文本
                            text = self._clean_ole_text(text)
                            if text.strip():
                                content += text + "\n"
                        except:
                            # 如果UTF-8失败，尝试其他编码
                            try:
                                text = data.decode('gbk', errors='ignore')
                                text = self._clean_ole_text(text)
                                if text.strip():
                                    content += text + "\n"
                            except:
                                pass
                        stream.close()
                    except:
                        continue
            
            return content
            
        except Exception as e:
            logger.warning(f"从OLE文件提取文本失败: {str(e)}")
            return ""
    
    def _clean_ole_text(self, text):
        """清理从OLE文件提取的文本"""
        # 移除控制字符
        import re
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        return text.strip() 